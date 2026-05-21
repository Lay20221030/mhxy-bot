# -*- coding: utf-8 -*-
"""生成完整的 Windows 安装部署教程 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ─── 页面设置 ───
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ─── 样式 ───
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

for level, (size, color, bold) in {
    1: (22, '1A56DB', True),
    2: (16, '2563EB', True),
    3: (13, '333333', True),
}.items():
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(*bytes.fromhex(color))
    hs.font.bold = bold
    hs.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    hs.paragraph_format.space_after = Pt(8)

# 代码样式
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
code_style.paragraph_format.left_indent = Cm(1.2)
code_style.paragraph_format.space_before = Pt(2)
code_style.paragraph_format.space_after = Pt(2)
code_style.paragraph_format.line_spacing = 1.2

def add_code(text):
    for line in text.strip().split('\n'):
        doc.add_paragraph(line, style='CodeBlock')

def add_note(text, icon='💡'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'{icon}  {text}')
    run.font.size = Pt(9.5)
    run.font.italic = True

def add_step(num, title, desc=None):
    p = doc.add_paragraph()
    run = p.add_run(f'步骤 {num}：{title}')
    run.bold = True
    run.font.size = Pt(11)
    if desc:
        doc.add_paragraph(desc)

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph('')

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('梦幻西游')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('全自动多账号脚本')
run2.font.size = Pt(24)
run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph('')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Windows 安装 · 部署 · 测试指南\n').font.size = Pt(14)
meta.add_run('\n版本：v2.0（23集教程全部集成）\n').font.size = Pt(11)
r3 = meta.add_run('GitHub：https://github.com/Lay20221030/mhxy-bot')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_paragraph('')
doc.add_paragraph('')
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run('2026年5月').font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════
# 目录
# ══════════════════════════════════════════════════════════════
doc.add_heading('目  录', level=1)

toc = [
    ('一', '环境要求', '硬件、系统、游戏客户端'),
    ('二', '安装 Python', '下载、安装、验证'),
    ('三', '下载项目', 'ZIP 下载 或 Git 克隆'),
    ('四', '安装依赖包', 'pip install + 国内镜像加速'),
    ('五', '配置游戏窗口位置', '最关键的步骤 —— 获取坐标 + 修改配置文件'),
    ('六', '设置壁纸对齐', '4K 桌面壁纸对齐 5 个窗口'),
    ('七', '制作模板截图', '8 个必须模板 + 截图方法'),
    ('八', '启动脚本', '6 种模式的启动命令'),
    ('九', '快捷键说明', 'F12 暂停、F11 退出'),
    ('十', '功能模块说明', '14 个模块总览 + 6 种任务模式'),
    ('十一', '安全测试流程', '从安全区到实战的渐进测试'),
    ('十二', '常见问题排查', '8 个 FAQ'),
    ('十三', '日志与调试', '日志文件位置 + 调试方法'),
    ('十四', 'Windows 特有问题', '权限、缩放、防火墙'),
    ('附录 A', '模板截图清单', '19 个模板详细说明'),
    ('附录 B', '项目文件结构', '完整目录树'),
]

for num, title, desc in toc:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}. {title}')
    run.bold = True
    if desc:
        p2 = doc.add_paragraph(f'     {desc}')
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)

page_break()

# ══════════════════════════════════════════════════════════════
# 一、环境要求
# ══════════════════════════════════════════════════════════════
doc.add_heading('一、环境要求', level=1)

add_table(
    ['项目', '要求', '说明'],
    [
        ['操作系统', 'Windows 10 / 11（64位）', '不支持 32 位系统'],
        ['Python 版本', '3.9 或更高', '推荐 Python 3.11+'],
        ['游戏客户端', '梦幻西游国服原版客户端', '需要同时运行 5 个窗口'],
        ['显示器分辨率', '4K (3840×2160) 推荐', '或能平铺 5 个窗口的任意分辨率'],
        ['磁盘空间', '约 2GB', '含 Python + 依赖 + PaddleOCR 模型'],
        ['内存', '8GB 以上', '5 个游戏客户端 + OCR 引擎'],
        ['网络', '稳定连接', '游戏需要在线，GitHub 需要下载'],
    ]
)

# ══════════════════════════════════════════════════════════════
# 二、安装 Python
# ══════════════════════════════════════════════════════════════
doc.add_heading('二、安装 Python', level=1)

doc.add_paragraph('1. 打开浏览器访问 https://www.python.org/downloads/')
doc.add_paragraph('2. 点击黄色 "Download Python" 按钮下载最新版本')
doc.add_paragraph('3. 双击安装包，在安装界面：')
for item in ['✅ 务必勾选底部的 "Add Python to PATH"（最重要的一步！）',
             '点击 "Install Now" 开始安装',
             '等待安装完成，点击 "Close"']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('4. 验证安装：按 Win + R，输入 cmd 回车，依次输入：')

add_code('''
python --version    （应显示 Python 3.x.x）
pip --version       （应显示 pip 2x.x.x）
''')
add_note('如果提示 "python 不是内部命令"，说明没勾选 Add to PATH，需重装。')

# ══════════════════════════════════════════════════════════════
# 三、下载项目
# ══════════════════════════════════════════════════════════════
doc.add_heading('三、下载项目', level=1)

doc.add_heading('方式一：直接下载 ZIP（推荐）', level=3)
for item in ['打开 https://github.com/Lay20221030/mhxy-bot',
             '点击绿色 "<> Code" 按钮 → 选择 "Download ZIP"',
             '解压到桌面（或任意目录，路径不要有中文）',
             '重命名文件夹为 mhxy-bot（去掉 -main 后缀）']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('方式二：Git 克隆', level=3)
add_code('''
git clone https://github.com/Lay20221030/mhxy-bot.git
cd mhxy-bot
''')

# ══════════════════════════════════════════════════════════════
# 四、安装依赖包
# ══════════════════════════════════════════════════════════════
doc.add_heading('四、安装依赖包', level=1)

doc.add_paragraph('打开命令提示符，进入项目目录：')
add_code('''
cd C:\\Users\\你的用户名\\Desktop\\mhxy-bot
''')
doc.add_paragraph('执行安装（二选一）：')
add_code('''
# 标准安装（从 PyPI 官方源）
pip install -r requirements.txt

# 国内镜像加速（如果速度慢）
pip install -i https://pypi.tuna.tsinghua.edu.cn/simple -r requirements.txt
''')

doc.add_paragraph('')
doc.add_paragraph('会安装的库：')

add_table(
    ['库名', '用途', '大小'],
    [
        ['mss', '快速屏幕截图（比 PIL 快 10 倍）', '~50KB'],
        ['opencv-python', '图像识别、模板匹配、颜色检测', '~50MB'],
        ['numpy', '数值计算', '~15MB'],
        ['pyautogui', '鼠标键盘模拟、屏幕坐标操作', '~100KB'],
        ['Pillow', '图像处理', '~3MB'],
        ['keyboard', '全局热键监听（F11/F12）', '~50KB'],
        ['PyQt5', 'UI 控制面板（可选）', '~50MB'],
        ['paddleocr', 'PaddleOCR 文字识别（可选）', '~200MB'],
    ]
)

add_note('paddleocr 如果安装失败不影响核心功能，脚本会自动回退到 Tesseract 或颜色检测。')

# ══════════════════════════════════════════════════════════════
# 五、配置游戏窗口位置
# ══════════════════════════════════════════════════════════════
doc.add_heading('五、配置游戏窗口位置（最关键）', level=1)

doc.add_paragraph('这一步决定脚本能否正确点击到游戏窗口。')
doc.add_paragraph('')

add_step(1, '排列游戏窗口')
for item in ['打开 5 个梦幻西游客户端',
             '将窗口在屏幕上平铺（建议 3×2 或 5×1 排列）',
             '每个窗口设置为相同大小（推荐 1282×1000）',
             '确保所有窗口完全可见，不被任务栏或其他窗口遮挡']:
    doc.add_paragraph(item, style='List Bullet')

add_step(2, '获取每个窗口的屏幕坐标')
doc.add_paragraph('打开 Python 交互环境（Win+R → 输入 python → 回车）：')
add_code('''
import pyautogui
print("把鼠标移到窗口1左上角，5秒后自动读取...")
import time; time.sleep(5)
print(pyautogui.position())
''')
doc.add_paragraph('依次将鼠标移到 5 个窗口的左上角，记录显示的 (x, y) 值。')

add_step(3, '修改配置文件')
doc.add_paragraph('用记事本打开项目目录中的 config\\settings.py，找到 ACCOUNT_WINDOWS：')
add_code('''
ACCOUNT_WINDOWS = [
    # 号1 (队长) - 第一个窗口
    {"name": "号1(队长)", "x": 0,     "y": 53,    "width": 1282, "height": 1000},
    # 号2 - 第二个窗口
    {"name": "号2",     "x": 1282,  "y": 53,    "width": 1282, "height": 1000},
    # 号3 - 第三个窗口
    {"name": "号3",     "x": 2564,  "y": 53,    "width": 1282, "height": 1000},
    # 号4 - 第四个窗口
    {"name": "号4",     "x": 0,     "y": 1106,  "width": 1282, "height": 1000},
    # 号5 - 第五个窗口
    {"name": "号5",     "x": 1282,  "y": 1106,  "width": 1282, "height": 1000},
]

SCREEN_RESOLUTION = (3840, 2160)   # 你的显示器实际分辨率
''')
doc.add_paragraph('把 x、y 替换为步骤 2 记录的坐标，width/height 替换为窗口的实际宽高。')

add_note('x 和 y 是窗口左上角在屏幕上的绝对像素坐标。如果不确定窗口大小，保持默认 1024×768，只改坐标。')

# ══════════════════════════════════════════════════════════════
# 六、设置壁纸对齐
# ══════════════════════════════════════════════════════════════
doc.add_heading('六、设置壁纸对齐（可选，强烈推荐）', level=1)

doc.add_paragraph('项目中包含一张专门的对齐壁纸：梦幻西游5开对齐壁纸_3840x2160.png')
doc.add_paragraph('')

for item in ['右键 → 设为桌面背景',
             '右击桌面 → 显示设置 → 缩放设为 "125%"',
             '确认分辨率为 3840×2160',
             '将 5 个游戏窗口分别拖入壁纸上的 5 个方框内']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('如果显示器不是 4K，可以重新生成适配壁纸：')
add_code('''
python generate_wallpaper.py   （修改脚本中的 WALL_W 和 WALL_H）
''')

# ══════════════════════════════════════════════════════════════
# 七、制作模板截图
# ══════════════════════════════════════════════════════════════
doc.add_heading('七、制作模板截图', level=1)

doc.add_paragraph('脚本通过 OpenCV 图像识别判断游戏状态，需要在 templates/ 目录下放置对应的 UI 截图。')

doc.add_paragraph('')
doc.add_paragraph('截图方法：')
for item in ['QQ 截图：Ctrl+Alt+A → 框选元素 → 保存到 templates/',
             '微信截图：Alt+A',
             'Windows 截图：Win+Shift+S']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('必须制作的模板（8 个）：')

add_table(
    ['模板文件名', '内容说明', '怎么截'],
    [
        ['dialog_confirm.png', '对话中的"确认/继续"按钮', '截取对话框底部按钮区域'],
        ['combat_enemy_area.png', '战斗界面中的敌人区域', '进入战斗后截取一个敌人头像'],
        ['combat_end_dialog.png', '战斗结束后的结算窗口', '截取结算弹窗的标题部分'],
        ['quest_npc_flag.png', 'NPC 头顶的金色任务标记', '截取NPC头顶的问号/叹号'],
        ['death_dialog.png', '角色死亡后的复活对话框', '截取死亡弹窗'],
        ['treasure_map_icon.png', '背包中藏宝图的图标', '打开背包截取藏宝图小图标'],
        ['captcha_dialog.png', '验证码弹出窗口', '遇到验证码时截图保存'],
        ['login_screen.png', '游戏登录界面的特征', '截取登录界面标识部分'],
    ]
)

add_note('模板图片尽量截取清晰、无遮挡的小区域（建议 30×30 到 100×100 像素）。')

# ══════════════════════════════════════════════════════════════
# 八、启动脚本
# ══════════════════════════════════════════════════════════════
doc.add_heading('八、启动脚本', level=1)

doc.add_paragraph('在命令提示符中进入项目目录，运行：')

add_code('''
python main.py quest        # 师门任务模式
python main.py ghost        # 捉鬼任务模式
python main.py treasure_map # 藏宝图挖掘模式
python main.py escort       # 押镖模式
python main.py dungeon      # 副本模式
python main.py story        # 主线任务模式
python main.py --ui         # 带 PyQt5 控制面板启动
''')

doc.add_paragraph('')
doc.add_paragraph('启动后脚本会自动：')
for item in ['初始化 5 个窗口的截图区域',
             '注册 F11（退出）、F12（暂停）全局热键',
             '进入主循环：检测验证码 → 检测死亡 → 执行任务 → 检测战斗',
             '日志实时输出到控制台和 logs/mhxy_bot.log 文件']:
    doc.add_paragraph(item, style='List Bullet')

# ══════════════════════════════════════════════════════════════
# 九、快捷键说明
# ══════════════════════════════════════════════════════════════
doc.add_heading('九、快捷键说明', level=1)

add_table(
    ['快捷键', '功能', '说明'],
    [
        ['F12', '暂停 / 恢复', '按一次暂停，再按一次恢复运行'],
        ['F11', '安全退出', '退出脚本主循环'],
        ['F9', '屏蔽玩家', '游戏内快捷键，屏蔽其他玩家显示'],
        ['Tab', '切换目标', '战斗中切换攻击目标'],
        ['F1~F8', '技能键', '对应技能槽 1~8'],
        ['M', '打开/关闭地图', '游戏内置地图快捷键'],
        ['I', '打开/关闭背包', '游戏内置背包快捷键'],
        ['J', '打开/关闭任务面板', '游戏内置任务面板快捷键'],
    ]
)

add_note('脚本运行期间请勿手动按 F1~F8 等技能键，会干扰脚本的自动战斗。')

# ══════════════════════════════════════════════════════════════
# 十、功能模块说明
# ══════════════════════════════════════════════════════════════
doc.add_heading('十、功能模块说明', level=1)

doc.add_paragraph('项目包含 14 个 Python 模块，支持 6 种游戏自动化模式：')

doc.add_heading('10.1  核心模块（core/）', level=3)
add_table(
    ['模块', '功能', '关键技术'],
    [
        ['window_group.py', '5 窗口管理', '坐标路由、ROI 截图、模板缓存'],
        ['screen.py', '屏幕捕获', 'mss 截图、OpenCV 颜色检测、模板查找'],
        ['input_sim.py', '输入模拟', '多窗口焦点切换、坐标转换、技能/药品'],
        ['captcha.py', '验证码处理', '滑块检测、缺口识别、随机抖动拖动模拟'],
        ['flow_control.py', '流控防封', '操作限频、随机延迟、滑动暂停'],
        ['session_manager.py', '会话管理', '掉线检测×3、自动重登、监控线程、换号'],
    ]
)

doc.add_heading('10.2  功能模块（modules/）', level=3)
add_table(
    ['模块', '功能', '关键技术'],
    [
        ['combat.py', '多号战斗', 'HP/MP 像素百分比检测、技能循环、自动吃药'],
        ['navigation.py', '导航系统', '三段导航（旗→驿站→步行）、遮挡穿透、位置缓存'],
        ['teleport.py', '飞行旗/飞行符', '红点检测、最近落点选择、地图列表定位'],
        ['station_coach.py', '驿站传送', 'F9 屏蔽干扰、移动 NPC 模糊匹配'],
        ['task_reader.py', '任务读取', '两阶段检测（颜色 0.1s → OpenCV 兜底）'],
        ['bandit_hunt.py', '贼王搜索', '野外坐标搜索 + 房间巡逻模式'],
        ['escort_landmark.py', '押镖地标', '地标匹配 + 卡屏 4 角哈希 + 战斗恢复'],
        ['warehouse.py', '仓库管理', '自动取图 + ALT 转图给其他号'],
        ['ocr_engine.py', 'OCR 文字识别', 'PaddleOCR → Tesseract → 颜色 三级回退'],
    ]
)

doc.add_heading('10.3  任务模式', level=3)
add_table(
    ['模式', '命令', '适合场景'],
    [
        ['师门任务', 'python main.py quest', '单人日常，挂机刷门派贡献'],
        ['捉鬼任务', 'python main.py ghost', '5 人组队，队长带队循环捉鬼'],
        ['藏宝图', 'python main.py treasure_map', '单人挖图，自动导航到挖掘点'],
        ['押镖', 'python main.py escort', '单人押镖，地标导航 + 战斗自动恢复'],
        ['副本', 'python main.py dungeon', '5 人组队副本清怪'],
        ['主线任务', 'python main.py story', '自动跟踪主线任务标记'],
    ]
)

# ══════════════════════════════════════════════════════════════
# 十一、安全测试流程
# ══════════════════════════════════════════════════════════════
doc.add_heading('十一、安全测试流程', level=1)

doc.add_paragraph('首次使用请按以下顺序逐步测试，不要直接在生产环境运行：')

for i, (phase, desc, items) in enumerate([
    ('第一阶段：坐标验证', '在安全区（长安城等无怪区域）测试',
     ['启动脚本 → 观察鼠标是否点击到了正确窗口',
      '按 F12 暂停 → 手动调整窗口位置 → 按 F12 恢复',
      '确认每个窗口都能被正确识别']),
    ('第二阶段：导航测试', '测试自动寻路',
     ['在安全区启动师门任务模式',
      '观察是否打开了地图并点击了任务标记',
      '确认角色自动寻路到达目标']),
    ('第三阶段：战斗测试', '在低等级怪物区域测试',
     ['启动脚本 → 走到有怪区域',
      '观察战斗检测是否触发',
      '确认 HP/MP 检测和技能释放是否正常']),
    ('第四阶段：正式使用', '确认前三阶段无误后',
     ['在你要挂机的场景启动对应模式',
      '建议先人工监督 10 分钟',
      '确认稳定后再无人值守']),
]):
    doc.add_heading(phase, level=3)
    doc.add_paragraph(desc)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

add_note('测试期间务必保持 F11/F12 快捷键可用，一旦出现异常立即暂停或退出。')

# ══════════════════════════════════════════════════════════════
# 十二、常见问题排查
# ══════════════════════════════════════════════════════════════
doc.add_heading('十二、常见问题排查', level=1)

qa = [
    ('Q1: 鼠标乱飞，点击位置不对？',
     '检查 config/settings.py 中 ACCOUNT_WINDOWS 的坐标。x 和 y 必须是窗口左上角在屏幕上的绝对像素坐标，而不是窗口内的相对坐标。用 pyautogui.position() 重新测量。'),
    ('Q2: 识别不到游戏元素？',
     '1) 确认 templates/ 目录下有对应的 PNG 模板截图\n2) 在 settings.py 中降低 MATCH_THRESHOLD（从 0.85 到 0.75）\n3) 确认模板截图和当前游戏画面风格一致（分辨率、UI 皮肤相同）'),
    ('Q3: 按键发到了别的窗口？',
     '脚本的 InputSim 会自动点击窗口切换焦点。确保 5 个游戏窗口都在屏幕上可见，没有被其他窗口遮挡。如果桌面有多个显示器，确保所有窗口在主显示器上。'),
    ('Q4: 脚本卡住了不动？',
     '脚本内置自动恢复机制：连续卡住 3 次会自动按 ESC 关闭弹窗。也可以按 F12 暂停 → 手动处理 → 按 F12 恢复。查看 logs/mhxy_bot.log 了解卡住原因。'),
    ('Q5: ImportError: No module named xxx？',
     '依赖没装全。重新运行：pip install -r requirements.txt。如果是 paddleocr 报错，可以跳过：pip install mss opencv-python numpy pyautogui Pillow keyboard PyQt5'),
    ('Q6: pyautogui 报错 PermissionError？',
     '以管理员身份运行命令提示符（右键 → 以管理员身份运行），再执行脚本。Windows 某些安全策略会阻止普通权限的鼠标键盘模拟。'),
    ('Q7: 验证码弹窗不会自动处理？',
     '滑块验证码脚本会自动检测缺口位置并模拟手动滑动。如果自动处理失败，脚本会暂停等待手动操作。按 F12 恢复继续。'),
    ('Q8: PaddleOCR 初始化失败？',
     'PaddleOCR 需要约 2GB 额外空间下载模型。如果安装失败，脚本会自动回退到 Tesseract OCR 或颜色检测模式，不影响核心功能。'),
]

for question, answer in qa:
    p = doc.add_paragraph()
    run = p.add_run(question)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    doc.add_paragraph(answer)

# ══════════════════════════════════════════════════════════════
# 十三、日志与调试
# ══════════════════════════════════════════════════════════════
doc.add_heading('十三、日志与调试', level=1)

doc.add_paragraph('脚本运行时会实时输出日志到两个位置：')
doc.add_paragraph('')

add_table(
    ['输出位置', '说明'],
    [
        ['控制台（cmd 窗口）', '实时显示运行状态，方便即时观察'],
        ['logs/mhxy_bot.log', '详细日志文件，用于事后排查'],
    ]
)

doc.add_paragraph('')
doc.add_paragraph('日志记录的内容包括：')
for item in ['每次操作的窗口序号和屏幕坐标',
             '模板匹配结果和置信度（如 "找到 quest_npc_flag 置信度 0.92"）',
             'HP/MP 检测百分比',
             '战斗状态变化',
             '错误和异常信息',
             '卡顿恢复操作记录']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('调试技巧：')
for item in ['遇到问题先查看日志最后 20 行（通常能找到原因）',
             '在 settings.py 中设置 logging.INFO → logging.DEBUG 可看到更多细节',
             '使用 --ui 参数启动可以看到可视化的日志面板']:
    doc.add_paragraph(item, style='List Bullet')

# ══════════════════════════════════════════════════════════════
# 十四、Windows 特有问题
# ══════════════════════════════════════════════════════════════
doc.add_heading('十四、Windows 特有问题', level=1)

doc.add_paragraph('Windows 系统特有的注意事项：')

add_table(
    ['问题', '原因', '解决方法'],
    [
        ['管理员权限', 'pyautogui 模拟输入需要管理员权限', '右键 cmd → 以管理员身份运行'],
        ['DPI 缩放', '高 DPI 显示器可能导致坐标偏移', '桌面右键 → 显示设置 → 缩放 125%'],
        ['防火墙拦截', 'pyautogui/keyboard 可能被安全软件拦截', '将 Python 加入防火墙白名单'],
        ['中文路径', '路径中包含中文可能导致模块加载失败', '项目放在纯英文路径下（如 C:\\mhxy-bot）'],
        ['窗口焦点', 'Windows 11 的虚拟桌面可能丢失焦点', '所有窗口放在同一桌面'],
        ['游戏全屏', '全屏模式无法多窗口同时操作', '游戏必须设为窗口模式'],
        ['多显示器', '多显示器排列可能导致坐标超出范围', '所有窗口放在主显示器上'],
    ]
)

page_break()

# ══════════════════════════════════════════════════════════════
# 附录 A：模板截图清单
# ══════════════════════════════════════════════════════════════
doc.add_heading('附录 A：模板截图清单', level=1)

doc.add_paragraph('以下 19 个模板用于图像识别，按重要程度排序：')

add_table(
    ['模板文件名', '用途', '重要程度', '截图区域'],
    [
        ['dialog_confirm.png', '对话框确认/继续按钮', '★必须', '对话框底部按钮'],
        ['combat_enemy_area.png', '战斗中敌人区域标识', '★必须', '任意敌人的头像'],
        ['combat_end_dialog.png', '战斗结束结算窗口', '★必须', '结算弹窗标题部分'],
        ['quest_npc_flag.png', 'NPC 头顶任务标记', '★必须', '金色问号/叹号'],
        ['quest_accept_btn.png', '接任务确认按钮', '★必须', '任务面板中的接受按钮'],
        ['death_dialog.png', '角色死亡复活对话框', '★必须', '死亡弹窗'],
        ['treasure_map_icon.png', '背包中的藏宝图图标', '☆建议', '背包里的宝图小图标'],
        ['captcha_dialog.png', '验证码弹窗', '☆建议', '验证码对话框'],
        ['login_screen.png', '登录界面标识', '☆建议', '登录界面的特征区域'],
        ['disconnect_dialog.png', '掉线提示弹窗', '☆建议', '掉线弹窗'],
        ['fly_flag_icon.png', '背包中的飞行旗', '☆建议', '背包里的飞行旗图标'],
        ['fly_scroll_icon.png', '背包中的飞行符', '☆建议', '背包里的飞行符图标'],
        ['backpack_open.png', '背包界面标识', '○可选', '背包打开后的特征区域'],
        ['warehouse_npc.png', '仓库管理员 NPC', '○可选', '仓库NPC身体部分'],
        ['friend_icon.png', '好友列表中的好友图标', '○可选', '好友列表中好友头像'],
        ['login_button.png', '登录界面登录按钮', '○可选', '登录按钮'],
        ['quest_submit_btn.png', '交任务按钮', '○可选', '任务面板中的提交按钮'],
        ['dialog_continue.png', '对话继续/下一步', '○可选', '对话框中的继续按钮'],
        ['dialog_close.png', '对话框关闭按钮', '○可选', '对话框右上角X'],
    ]
)

add_note('★必须 = 核心功能依赖  ☆建议 = 增强稳定性  ○可选 = 特定任务需要')

page_break()

# ══════════════════════════════════════════════════════════════
# 附录 B：项目文件结构
# ══════════════════════════════════════════════════════════════
doc.add_heading('附录 B：项目文件结构', level=1)

add_code('''
mhxy-bot/
├── main.py                     # 主控入口（AutoBot 类，6 种模式）
├── ui.py                       # PyQt5 控制面板
├── requirements.txt            # Python 依赖清单
├── .gitignore                  # 忽略 __pycache__ / logs / templates
│
├── config/
│   └── settings.py             # 全部配置（窗口/导航/战斗/传送/模板）
│
├── core/                       # 核心基础设施
│   ├── window_group.py         #   5 窗口管理器（GameWindow + WindowGroup）
│   ├── screen.py               #   屏幕捕获（mss + OpenCV 颜色/模板）
│   ├── input_sim.py            #   输入模拟（多窗口路由 + 焦点切换）
│   ├── captcha.py              #   验证码处理（滑块/图形/对话框）
│   ├── flow_control.py         #   流控防封（限频/延迟/暂停）
│   └── session_manager.py     #   会话管理（掉线/重登/换号）
│
├── modules/                    # 功能模块
│   ├── combat.py               #   多账号战斗（HP/MP检测 + 技能循环）
│   ├── navigation.py           #   三段导航（飞行旗→驿站→步行）
│   ├── teleport.py             #   飞行旗/飞行符传送
│   ├── station_coach.py        #   驿站车夫传送
│   ├── task_reader.py          #   两阶段任务读取
│   ├── bandit_hunt.py          #   强盗/贼王搜索
│   ├── escort_landmark.py     #   押镖地标检测 + 卡屏判定
│   ├── warehouse.py            #   仓库取图 + ALT 转图
│   ├── ocr_engine.py           #   PaddleOCR 文字识别引擎
│   └── tasks/                  #   任务处理器
│       ├── sect_quest.py       #     师门任务（4 种类型）
│       ├── ghost_hunt.py       #     捉鬼任务（多轮循环）
│       └── treasure_map.py    #     藏宝图挖掘
│
├── templates/                  # 模板截图目录（需要自行制作）
│   └── *.png                   #   至少准备 6~8 个 PNG 模板
│
└── logs/                       # 运行日志
    └── mhxy_bot.log            #   自动生成
''')

# ══════════════════════════════════════════════════════════════
# 页脚
# ══════════════════════════════════════════════════════════════
page_break()

for _ in range(10):
    doc.add_paragraph('')

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run('— 文档结束 —')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph('')

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
ir = info.add_run('GitHub: https://github.com/Lay20221030/mhxy-bot\n版本: v2.0 | 2026年5月')
ir.font.size = Pt(9)
ir.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ══════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════
output = '/Users/lay/Desktop/梦幻西游自动脚本_Windows安装部署教程.docx'
doc.save(output)
size_kb = os.path.getsize(output) / 1024
print(f'文档已保存: {output}')
print(f'文件大小: {size_kb:.1f} KB')
