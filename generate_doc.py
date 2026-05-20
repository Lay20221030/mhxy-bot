# -*- coding: utf-8 -*-
"""生成 Windows 安装部署教程 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ─── 样式设置 ───
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# 标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = '微软雅黑'
    heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB) if i == 1 else RGBColor(0x33, 0x33, 0x33)

# 代码块样式
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Cm(1)

def add_code_block(text):
    """添加代码块"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line, style='CodeBlock')

def add_note(text):
    """添加提示框"""
    p = doc.add_paragraph()
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x85, 0x6E, 0x0B)

def add_step(num, title, desc=""):
    """添加步骤"""
    p = doc.add_paragraph()
    run = p.add_run(f'步骤 {num}：{title}')
    run.bold = True
    run.font.size = Pt(12)
    if desc:
        doc.add_paragraph(desc)

# ════════════════════════════════════════
# 封面
# ════════════════════════════════════════
title = doc.add_heading('梦幻西游自动脚本', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Windows 安装部署教程')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('版本：v1.0').font.size = Pt(11)
info.add_run('\nGitHub：https://github.com/Lay20221030/mhxy-bot').font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════
# 目录（简化版）
# ════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '1. 环境要求',
    '2. 安装 Python',
    '3. 下载项目代码',
    '4. 安装依赖包',
    '5. 配置游戏窗口位置',
    '6. 制作模板截图',
    '7. 启动脚本',
    '8. 快捷键说明',
    '9. 模式说明',
    '10. 常见问题排查',
    '11. 日志查看',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════
# 1. 环境要求
# ════════════════════════════════════════
doc.add_heading('1. 环境要求', level=1)

doc.add_paragraph('操作系统：Windows 10 / Windows 11（64位）')
doc.add_paragraph('Python 版本：3.9 或更高')
doc.add_paragraph('游戏客户端：梦幻西游国服原版客户端')
doc.add_paragraph('屏幕分辨率：能平铺 5 个游戏窗口')
doc.add_paragraph('磁盘空间：约 500MB（Python + 依赖）')

# ════════════════════════════════════════
# 2. 安装 Python
# ════════════════════════════════════════
doc.add_heading('2. 安装 Python', level=1)

doc.add_paragraph('2.1 下载 Python')
doc.add_paragraph('打开浏览器，访问 https://www.python.org/downloads/')
doc.add_paragraph('点击黄色的 "Download Python" 按钮，下载最新版本。')

doc.add_paragraph('2.2 安装 Python')
doc.add_paragraph('双击下载的安装包，在安装界面：')
items = [
    '务必勾选底部的 "Add Python to PATH"（非常重要！）',
    '点击 "Install Now" 开始安装',
    '等待安装完成，点击 "Close"',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_note('如果忘记勾选 Add Python to PATH，需要卸载后重新安装，或手动添加环境变量。')

doc.add_paragraph('2.3 验证安装')
doc.add_paragraph('按 Win + R，输入 cmd 回车，打开命令提示符。依次输入以下命令验证：')

add_code_block('''
python --version
pip --version
''')

doc.add_paragraph('如果显示版本号，说明安装成功。')

# ════════════════════════════════════════
# 3. 下载项目代码
# ════════════════════════════════════════
doc.add_heading('3. 下载项目代码', level=1)

doc.add_paragraph('方式一：直接下载 ZIP（推荐）')
items = [
    '打开浏览器访问：https://github.com/Lay20221030/mhxy-bot',
    '点击绿色的 "Code" 按钮',
    '选择 "Download ZIP"',
    '解压到桌面或任意目录',
    '重命名文件夹为 mhxy-bot',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('方式二：Git 克隆（适合有 Git 的用户）')

add_code_block('''
git clone https://github.com/Lay20221030/mhxy-bot.git
cd mhxy-bot
''')

# ════════════════════════════════════════
# 4. 安装依赖包
# ════════════════════════════════════════
doc.add_heading('4. 安装依赖包', level=1)

doc.add_paragraph('打开命令提示符（Win + R → cmd → 回车），进入项目目录：')

add_code_block('''
cd C:\\Users\\你的用户名\\Desktop\\mhxy-bot
''')

doc.add_paragraph('执行安装命令：')

add_code_block('''
pip install -r requirements.txt
''')

doc.add_paragraph('等待安装完成，会依次安装以下库：')
items = [
    'mss — 快速屏幕截图',
    'opencv-python — 图像识别',
    'numpy — 数值计算',
    'pyautogui — 鼠标键盘模拟',
    'Pillow — 图像处理',
    'keyboard — 键盘热键监听',
    'PyQt5 — UI 面板（可选）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_note('如果安装速度慢，可以使用国内镜像：pip install -i https://pypi.tuna.tsinghua.edu.cn/simple -r requirements.txt')

# ════════════════════════════════════════
# 5. 配置游戏窗口位置
# ════════════════════════════════════════
doc.add_heading('5. 配置游戏窗口位置', level=1)

doc.add_paragraph('这是最关键的一步，需要精确配置每个游戏窗口的屏幕坐标。')

doc.add_paragraph('5.1 排列游戏窗口')
items = [
    '打开 5 个梦幻西游客户端',
    '将 5 个窗口在屏幕上水平并排铺开',
    '建议每个窗口设置为相同大小（如 1024×768）',
    '确保所有窗口完全可见，不要被遮挡',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('5.2 获取窗口坐标')
doc.add_paragraph('方法一：使用 Python 脚本获取')

add_code_block('''
# 在命令提示符中运行：
python -c "import pyautogui; print('当前鼠标位置:', pyautogui.position())"
''')

doc.add_paragraph('将鼠标移动到每个窗口的左上角，记录显示的坐标。')
doc.add_paragraph('方法二：使用截图工具')
doc.add_paragraph('按 Win + Shift + S 截图，用画图工具打开，将鼠标移到窗口左上角，记录左下角显示的像素坐标。')

doc.add_paragraph('5.3 修改配置文件')
doc.add_paragraph('用记事本打开项目目录中的 config\\settings.py 文件，找到 ACCOUNT_WINDOWS 部分并修改：')

add_code_block('''
ACCOUNT_WINDOWS = [
    # 号1 (队长) - 第一个窗口的左上角坐标和尺寸
    {"name": "号1(队长)", "x": 0,     "y": 0,     "width": 1024, "height": 768},
    # 号2
    {"name": "号2",     "x": 1024,  "y": 0,     "width": 1024, "height": 768},
    # 号3
    {"name": "号3",     "x": 2048,  "y": 0,     "width": 1024, "height": 768},
    # 号4
    {"name": "号4",     "x": 3072,  "y": 0,     "width": 1024, "height": 768},
    # 号5
    {"name": "号5",     "x": 4096,  "y": 0,     "width": 1024, "height": 768},
]
''')

doc.add_paragraph('根据实际测量值修改每个窗口的 x、y、width、height。')
doc.add_paragraph('同时修改 SCREEN_RESOLUTION 为你的显示器总分辨率：')

add_code_block('''
# 5个1024宽的窗口 = 5120，高度 = 768
SCREEN_RESOLUTION = (5120, 768)
''')

add_note('x 和 y 是窗口左上角在屏幕上的像素坐标。width 和 height 是窗口的宽高。如果不确定，保持默认尺寸 1024×768，只改 x 坐标。')

# ════════════════════════════════════════
# 6. 制作模板截图
# ════════════════════════════════════════
doc.add_heading('6. 制作模板截图', level=1)

doc.add_paragraph('脚本通过图像识别来判断游戏状态，需要预先截取一些 UI 元素的图片作为模板。')

doc.add_paragraph('6.1 用 QQ/微信截图工具截图')
items = [
    '在游戏中找到对应元素',
    '按 Ctrl+Alt+A（QQ）或 Alt+A（微信）截图',
    '框选元素区域，保存为 PNG 格式',
    '保存到项目目录的 templates 文件夹中',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('6.2 需要的模板截图')

# 模板表格
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'

# 表头
headers = ['文件名', '内容说明', '重要程度']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

# 表格内容
rows_data = [
    ('quest_npc_flag.png', '任务 NPC 头顶的金色问号/叹号标记', '必须'),
    ('dialog_confirm.png', '对话中的确认/继续按钮', '必须'),
    ('combat_enemy_area.png', '战斗中敌方怪物的区域', '必须'),
    ('combat_end_dialog.png', '战斗结束后的结算对话框', '必须'),
    ('death_dialog.png', '角色死亡后的复活对话框', '建议'),
    ('treasure_map_icon.png', '背包中的藏宝图图标', '可选'),
    ('captcha_dialog.png', '验证码弹出对话框', '建议'),
    ('dialog_continue.png', '对话中的继续/下一步按钮', '建议'),
]
for i, (fname, desc, importance) in enumerate(rows_data):
    table.rows[i+1].cells[0].text = fname
    table.rows[i+1].cells[1].text = desc
    table.rows[i+1].cells[2].text = importance

doc.add_paragraph('')

add_note('模板截图的质量直接影响识别成功率。尽量截取清晰、无干扰的图片，尺寸不要太大（建议 30×30 ~ 100×100 像素）。')

doc.add_paragraph('6.3 调整识别阈值')
doc.add_paragraph('如果识别不稳定，可以在 config\\settings.py 中调整：')

add_code_block('''
# 图像匹配阈值（0~1），越低越宽松
MATCH_THRESHOLD = 0.80   # 默认 0.85，可降到 0.80
SCREEN_SCALE = 0.5       # 缩放比例，越大越精确但越慢
''')

# ════════════════════════════════════════
# 7. 启动脚本
# ════════════════════════════════════════
doc.add_heading('7. 启动脚本', level=1)

doc.add_paragraph('在命令提示符中进入项目目录，运行以下命令：')

add_code_block('''
# 师门任务模式
python main.py quest

# 捉鬼模式
python main.py ghost

# 藏宝图模式
python main.py treasure_map

# 押镖模式
python main.py escort

# 副本模式
python main.py dungeon

# 主线任务模式
python main.py story

# 带 UI 面板启动
python main.py --ui
''')

doc.add_paragraph('启动后，脚本会自动：')
items = [
    '检测每个窗口的游戏状态',
    '队长接任务 → 队员确认',
    '导航到目标位置',
    '进入战斗后自动释放技能、吃药',
    '战斗结束后自动处理奖励',
    '循环执行直到手动停止',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_note('首次运行建议先在安全区域测试，确保坐标配置正确后再正式使用。')

# ════════════════════════════════════════
# 8. 快捷键说明
# ════════════════════════════════════════
doc.add_heading('8. 快捷键说明', level=1)

hotkey_table = doc.add_table(rows=5, cols=3)
hotkey_table.style = 'Light Grid Accent 1'

hotkey_headers = ['快捷键', '功能', '说明']
for i, h in enumerate(hotkey_headers):
    cell = hotkey_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

hotkey_data = [
    ('F12', '暂停/恢复', '按一次暂停，再按一次恢复运行'),
    ('F11', '退出脚本', '安全退出脚本'),
    ('F9', '自动战斗', '游戏中开启/关闭自动战斗'),
    ('Alt+Tab', '切换窗口', '手动切换到游戏窗口（不要用于脚本）'),
]
for i, (key, func, desc) in enumerate(hotkey_data):
    hotkey_table.rows[i+1].cells[0].text = key
    hotkey_table.rows[i+1].cells[1].text = func
    hotkey_table.rows[i+1].cells[2].text = desc

doc.add_paragraph('')
add_note('脚本运行期间不要按 F1~F8 等技能键，会触发游戏操作。')

# ════════════════════════════════════════
# 9. 模式说明
# ════════════════════════════════════════
doc.add_heading('9. 模式说明', level=1)

mode_table = doc.add_table(rows=7, cols=3)
mode_table.style = 'Light Grid Accent 1'

mode_headers = ['模式', '命令', '说明']
for i, h in enumerate(mode_headers):
    cell = mode_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

mode_data = [
    ('师门任务', 'python main.py quest', '自动接师门任务、完成、交任务，循环执行'),
    ('捉鬼任务', 'python main.py ghost', '队长带队捉鬼，自动识别鬼怪进入战斗'),
    ('藏宝图', 'python main.py treasure_map', '自动使用藏宝图，导航到挖掘点，领取奖励'),
    ('押镖', 'python main.py escort', '自动导航到镖师，护送途中战斗'),
    ('副本', 'python main.py dungeon', '自动进入副本，清怪'),
    ('主线任务', 'python main.py story', '自动跟踪主线任务标记'),
]
for i, (name, cmd, desc) in enumerate(mode_data):
    mode_table.rows[i+1].cells[0].text = name
    mode_table.rows[i+1].cells[1].text = cmd
    mode_table.rows[i+1].cells[2].text = desc

# ════════════════════════════════════════
# 10. 常见问题排查
# ════════════════════════════════════════
doc.add_heading('10. 常见问题排查', level=1)

qa = [
    ('Q: 鼠标乱飞，点击位置不对？',
     '检查 config/settings.py 中 ACCOUNT_WINDOWS 的坐标是否正确。确保 x、y 是窗口左上角在屏幕上的绝对坐标。'),
    ('Q: 识别不到游戏元素？',
     '1) 检查 templates/ 目录下是否有对应的模板截图\n2) 降低 MATCH_THRESHOLD 到 0.80\n3) 确认模板截图和游戏画面风格一致'),
    ('Q: 按键发到了别的窗口？',
     '脚本会自动点击窗口焦点。确保 5 个游戏窗口都在屏幕上可见，没有被其他窗口遮挡。'),
    ('Q: 脚本卡住了不动？',
     '脚本有自动恢复机制：连续卡住 3 次会自动按 ESC 关闭弹窗。也可以按 F12 暂停后手动处理，再按 F12 恢复。'),
    ('Q: ImportError: No module named xxx？',
     '说明依赖没装全，重新运行：pip install -r requirements.txt'),
    ('Q: pyautogui 权限不足？',
     '以管理员身份运行命令提示符（右键 → 以管理员身份运行），再执行脚本。'),
    ('Q: 验证码弹窗不会处理？',
     '滑块验证码会自动检测并尝试处理，失败后会暂停等待手动操作。按 F12 恢复。'),
]

for question, answer in qa:
    p = doc.add_paragraph()
    run = p.add_run(question)
    run.bold = True
    doc.add_paragraph(answer)

# ════════════════════════════════════════
# 11. 日志查看
# ════════════════════════════════════════
doc.add_heading('11. 日志查看', level=1)

doc.add_paragraph('脚本运行时会生成日志文件，方便排查问题：')

add_code_block('''
logs/mhxy_bot.log   # 详细运行日志
''')

doc.add_paragraph('日志记录内容包括：')
items = [
    '每次操作的窗口和坐标',
    '模板匹配结果和置信度',
    '战斗状态变化',
    '错误和异常信息',
    '卡顿恢复记录',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_note('遇到问题时，先查看日志文件末尾的内容，通常能找到错误原因。')

# ════════════════════════════════════════
# 页脚
# ════════════════════════════════════════
doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub: https://github.com/Lay20221030/mhxy-bot')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ─── 保存文档 ───
output_path = '/Users/lay/Desktop/梦幻西游自动脚本-Windows安装部署教程.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
